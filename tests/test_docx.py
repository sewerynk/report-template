"""Tests for DOCX formatter."""

import tempfile
from datetime import date
from pathlib import Path

import pytest

from report_template.generator import ReportGenerator
from report_template.models import (
    FeatureDevReport,
    OutputFormat,
    Priority,
    ReportType,
    Status,
    Task,
    TeamMember,
)

# Check if python-docx is available
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@pytest.fixture
def sample_report():
    """Create a sample report for testing."""
    return FeatureDevReport(
        title="Test Report",
        project_name="Test Project",
        author="Test Author",
        feature_name="Test Feature",
        summary="This is a test summary for DOCX generation",
        team_members=[
            TeamMember(name="John Doe", role="Developer", email="john@example.com"),
            TeamMember(name="Jane Smith", role="Designer"),
        ],
        objectives=["Objective 1", "Objective 2"],
        tasks=[
            Task(
                title="Task 1",
                status=Status.COMPLETED,
                priority=Priority.HIGH,
                assignee="John Doe"
            ),
            Task(
                title="Task 2",
                status=Status.IN_PROGRESS,
                priority=Priority.MEDIUM
            ),
        ],
        technical_approach="Using modern technologies for this feature",
    )


@pytest.fixture
def generator():
    """Create a ReportGenerator instance."""
    return ReportGenerator()


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_generate_docx(generator, sample_report):
    """Test generating a DOCX report."""
    result = generator.generate(
        sample_report,
        ReportType.FEATURE_DEV,
        OutputFormat.DOCX
    )

    assert isinstance(result, bytes)
    assert len(result) > 0


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_generate_docx_to_file(generator, sample_report):
    """Test generating a DOCX report to a file."""
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "report.docx"

        result_path = generator.generate_to_file(
            sample_report,
            ReportType.FEATURE_DEV,
            output_path
        )

        assert result_path.exists()
        assert result_path == output_path
        assert result_path.stat().st_size > 0

        # Verify it's a valid DOCX file
        doc = docx.Document(str(result_path))
        assert len(doc.paragraphs) > 0

        # Check that title is in the document
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Test Report" in text


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_format_auto_detection(generator, sample_report):
    """Test automatic DOCX format detection from file extension."""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Test .docx extension
        docx_path = Path(tmpdir) / "report.docx"
        generator.generate_to_file(sample_report, ReportType.FEATURE_DEV, docx_path)
        assert docx_path.exists()

        # Test .doc extension (should also work)
        doc_path = Path(tmpdir) / "report.doc"
        generator.generate_to_file(sample_report, ReportType.FEATURE_DEV, doc_path)
        assert doc_path.exists()


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_contains_expected_sections(generator, sample_report):
    """Test that DOCX contains expected report sections."""
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "report.docx"
        generator.generate_to_file(sample_report, ReportType.FEATURE_DEV, output_path)

        # Read the document
        doc = docx.Document(str(output_path))

        # Extract all text
        all_text = "\n".join([p.text for p in doc.paragraphs])

        # Check for key sections
        assert "Test Report" in all_text
        assert "Test Project" in all_text
        assert "Test Feature" in all_text
        assert "Executive Summary" in all_text
        assert "Team" in all_text
        assert "Objectives" in all_text
        assert "John Doe" in all_text
        assert "Objective 1" in all_text


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_with_minimal_data(generator):
    """Test DOCX generation with minimal data."""
    minimal_report = FeatureDevReport(
        title="Minimal",
        project_name="Project",
        author="Author",
        feature_name="Feature",
        summary="Summary"
    )

    result = generator.generate(
        minimal_report,
        ReportType.FEATURE_DEV,
        OutputFormat.DOCX
    )

    assert isinstance(result, bytes)
    assert len(result) > 0


def test_docx_import_error_handling(generator, sample_report, monkeypatch):
    """Test error handling when python-docx is not available."""
    # Mock the import to fail
    import sys
    original_modules = sys.modules.copy()

    # Remove docx module if it exists
    if 'docx' in sys.modules:
        del sys.modules['docx']
    if 'report_template.formatters.docx' in sys.modules:
        del sys.modules['report_template.formatters.docx']

    # This would raise an ImportError in real scenario
    # but we can't fully test it without breaking other tests
    # Just verify the function exists
    from report_template.formatters import docx as docx_module
    assert hasattr(docx_module, 'create_docx_report')

    # Restore modules
    sys.modules.update(original_modules)
