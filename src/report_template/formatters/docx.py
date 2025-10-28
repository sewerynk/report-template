"""
DOCX formatter utilities using python-docx.
"""

from datetime import date
from typing import Any, Dict, List, Union

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


def create_docx_report(data: Dict[str, Any], report_type: str) -> bytes:
    """
    Create a DOCX report from data.

    Args:
        data: Report data dictionary.
        report_type: Type of report (feature_dev, program_mgmt, engineering_init).

    Returns:
        DOCX file content as bytes.

    Raises:
        ImportError: If python-docx is not installed.
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX generation. "
            "Install it with: pip install python-docx"
        )

    doc = Document()

    # Set up styles
    _setup_styles(doc)

    # Generate report based on type
    if report_type == "feature_dev":
        _create_feature_dev_report(doc, data)
    elif report_type == "program_mgmt":
        _create_program_mgmt_report(doc, data)
    elif report_type == "engineering_init":
        _create_engineering_init_report(doc, data)
    else:
        raise ValueError(f"Unknown report type: {report_type}")

    # Save to bytes
    from io import BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()


def _setup_styles(doc: "Document") -> None:
    """Set up custom styles for the document."""
    styles = doc.styles

    # Heading styles are already defined, just customize colors if needed
    # Title style
    if 'Title' in styles:
        title_style = styles['Title']
        title_font = title_style.font
        title_font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray


def _add_metadata_section(doc: "Document", data: Dict[str, Any]) -> None:
    """Add metadata section to document."""
    # Add a light gray background table for metadata
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Shading Accent 1'

    metadata_fields = [
        ("Project", data.get("project_name", "N/A")),
        ("Author", data.get("author", "N/A")),
        ("Date", _format_date(data.get("date"))),
        ("Version", data.get("version", "1.0")),
    ]

    # Add report-specific fields
    if "feature_name" in data:
        metadata_fields.insert(1, ("Feature", data["feature_name"]))
    if "program_name" in data:
        metadata_fields.insert(1, ("Program", data["program_name"]))
    if "initiative_name" in data:
        metadata_fields.insert(1, ("Initiative", data["initiative_name"]))

    if "repository" in data and data["repository"]:
        metadata_fields.append(("Repository", data["repository"]))
    if "branch" in data and data["branch"]:
        metadata_fields.append(("Branch", data["branch"]))
    if "sprint" in data and data["sprint"]:
        metadata_fields.append(("Sprint", data["sprint"]))
    if "status" in data and data["status"]:
        status_value = data["status"]["value"] if isinstance(data["status"], dict) else str(data["status"])
        metadata_fields.append(("Status", status_value))

    for label, value in metadata_fields:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)

    doc.add_paragraph()  # Add spacing


def _add_table(doc: "Document", headers: List[str], rows: List[List[str]]) -> None:
    """Add a formatted table to the document."""
    if not rows:
        doc.add_paragraph("No data available.", style='Normal')
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    doc.add_paragraph()  # Add spacing


def _add_list(doc: "Document", items: List[str], style: str = 'List Bullet') -> None:
    """Add a bulleted or numbered list to the document."""
    if not items:
        doc.add_paragraph("None specified.", style='Normal')
        return

    for item in items:
        doc.add_paragraph(str(item), style=style)

    doc.add_paragraph()  # Add spacing


def _format_date(date_value: Any) -> str:
    """Format a date value to string."""
    if isinstance(date_value, date):
        return date_value.strftime("%Y-%m-%d")
    elif isinstance(date_value, str):
        return date_value
    return "N/A"


def _create_feature_dev_report(doc: "Document", data: Dict[str, Any]) -> None:
    """Create a feature development report."""
    # Title
    title = doc.add_heading(data.get("title", "Feature Development Report"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    _add_metadata_section(doc, data)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data.get("summary", "No summary provided."))
    doc.add_paragraph()

    # Team
    doc.add_heading("Team", level=1)
    team_members = data.get("team_members", [])
    if team_members:
        team_rows = []
        for member in team_members:
            team_rows.append([
                member.get("name", "N/A"),
                member.get("role", "N/A"),
                member.get("email", "N/A")
            ])
        _add_table(doc, ["Name", "Role", "Email"], team_rows)
    else:
        doc.add_paragraph("No team members specified.")
        doc.add_paragraph()

    # Objectives
    doc.add_heading("Objectives", level=1)
    _add_list(doc, data.get("objectives", []))

    # Requirements
    doc.add_heading("Requirements", level=1)
    _add_list(doc, data.get("requirements", []))

    # Technical Approach
    doc.add_heading("Technical Approach", level=1)
    doc.add_paragraph(data.get("technical_approach", "To be determined."))
    doc.add_paragraph()

    if data.get("architecture_notes"):
        doc.add_heading("Architecture Notes", level=2)
        doc.add_paragraph(data["architecture_notes"])
        doc.add_paragraph()

    # Tasks & Progress
    doc.add_heading("Tasks & Progress", level=1)
    tasks = data.get("tasks", [])
    if tasks:
        task_rows = []
        for task in tasks:
            status = task.get("status", {})
            priority = task.get("priority", {})
            status_value = status.get("value", "N/A") if isinstance(status, dict) else str(status)
            priority_value = priority.get("value", "N/A") if isinstance(priority, dict) else str(priority)

            task_rows.append([
                task.get("title", "N/A"),
                task.get("assignee", "Unassigned"),
                status_value,
                priority_value,
                task.get("jira_id", "N/A"),
                _format_date(task.get("target_start_date")),
                _format_date(task.get("target_end_date")),
                _format_date(task.get("due_date"))
            ])
        _add_table(doc, ["Task", "Assignee", "Status", "Priority", "JIRA", "Start Date", "End Date", "Due Date"], task_rows)
    else:
        doc.add_paragraph("No tasks defined.")
        doc.add_paragraph()

    if data.get("progress_notes"):
        doc.add_heading("Progress Notes", level=2)
        doc.add_paragraph(data["progress_notes"])
        doc.add_paragraph()

    # Milestones
    doc.add_heading("Milestones", level=1)
    milestones = data.get("milestones", [])
    if milestones:
        milestone_rows = []
        for milestone in milestones:
            status = milestone.get("status", {})
            status_value = status.get("value", "N/A") if isinstance(status, dict) else str(status)

            milestone_rows.append([
                milestone.get("name", "N/A"),
                _format_date(milestone.get("target_date")),
                status_value,
                f"{milestone.get('completion_percentage', 0)}%"
            ])
        _add_table(doc, ["Milestone", "Target Date", "Status", "Completion"], milestone_rows)
    else:
        doc.add_paragraph("No milestones defined.")
        doc.add_paragraph()

    # Testing Strategy
    doc.add_heading("Testing Strategy", level=1)
    doc.add_paragraph(data.get("testing_strategy", "To be defined."))
    doc.add_paragraph()

    # Deployment Plan
    doc.add_heading("Deployment Plan", level=1)
    doc.add_paragraph(data.get("deployment_plan", "To be defined."))
    doc.add_paragraph()

    # Risks & Mitigation
    doc.add_heading("Risks & Mitigation", level=1)
    risks = data.get("risks", [])
    if risks:
        for i, risk in enumerate(risks, 1):
            doc.add_heading(f"{i}. {risk.get('description', 'Risk')}", level=2)
            impact = risk.get("impact", {})
            likelihood = risk.get("likelihood", {})
            impact_value = impact.get("value", "N/A") if isinstance(impact, dict) else str(impact)
            likelihood_value = likelihood.get("value", "N/A") if isinstance(likelihood, dict) else str(likelihood)

            p = doc.add_paragraph()
            p.add_run("Impact: ").bold = True
            p.add_run(impact_value + "\n")
            p.add_run("Likelihood: ").bold = True
            p.add_run(likelihood_value + "\n")
            p.add_run("Mitigation: ").bold = True
            p.add_run(risk.get("mitigation", "N/A"))
            if risk.get("owner"):
                p.add_run("\nOwner: ").bold = True
                p.add_run(risk["owner"])
            doc.add_paragraph()
    else:
        doc.add_paragraph("No risks identified.")
        doc.add_paragraph()

    # Dependencies
    doc.add_heading("Dependencies", level=1)
    _add_list(doc, data.get("dependencies", []))

    # Footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on {_format_date(data.get('date', date.today()))}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _create_program_mgmt_report(doc: "Document", data: Dict[str, Any]) -> None:
    """Create a program management report."""
    # Title
    title = doc.add_heading(data.get("title", "Program Management Report"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    _add_metadata_section(doc, data)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data.get("executive_summary", data.get("summary", "No summary provided.")))
    doc.add_paragraph()

    # Stakeholders
    doc.add_heading("Stakeholders", level=1)
    _add_list(doc, data.get("stakeholders", []))

    # Team
    doc.add_heading("Team Composition", level=1)
    team_members = data.get("team_members", [])
    if team_members:
        team_rows = []
        for member in team_members:
            team_rows.append([
                member.get("name", "N/A"),
                member.get("role", "N/A"),
                member.get("email", "N/A")
            ])
        _add_table(doc, ["Name", "Role", "Email"], team_rows)
    else:
        doc.add_paragraph("No team members specified.")
        doc.add_paragraph()

    # KPIs
    doc.add_heading("Key Performance Indicators (KPIs)", level=1)
    kpis = data.get("kpis", {})
    if kpis:
        kpi_rows = [[k, str(v)] for k, v in kpis.items()]
        _add_table(doc, ["Metric", "Value"], kpi_rows)
    else:
        doc.add_paragraph("No KPIs defined.")
        doc.add_paragraph()

    # Milestones
    doc.add_heading("Milestones & Progress", level=1)
    milestones = data.get("milestones", [])
    if milestones:
        milestone_rows = []
        for milestone in milestones:
            status = milestone.get("status", {})
            status_value = status.get("value", "N/A") if isinstance(status, dict) else str(status)

            milestone_rows.append([
                milestone.get("name", "N/A"),
                _format_date(milestone.get("target_date")),
                status_value,
                f"{milestone.get('completion_percentage', 0)}%"
            ])
        _add_table(doc, ["Milestone", "Target Date", "Status", "Completion"], milestone_rows)
    else:
        doc.add_paragraph("No milestones defined.")
        doc.add_paragraph()

    # Key Achievements
    doc.add_heading("Key Achievements", level=1)
    _add_list(doc, data.get("key_achievements", []))

    # Challenges
    doc.add_heading("Challenges", level=1)
    _add_list(doc, data.get("challenges", []))

    # Risks
    doc.add_heading("Risks", level=1)
    risks = data.get("risks", [])
    if risks:
        risk_rows = []
        for risk in risks:
            impact = risk.get("impact", {})
            likelihood = risk.get("likelihood", {})
            impact_value = impact.get("value", "N/A") if isinstance(impact, dict) else str(impact)
            likelihood_value = likelihood.get("value", "N/A") if isinstance(likelihood, dict) else str(likelihood)

            risk_rows.append([
                risk.get("description", "N/A"),
                impact_value,
                likelihood_value,
                risk.get("mitigation", "N/A"),
                risk.get("owner", "N/A")
            ])
        _add_table(doc, ["Risk", "Impact", "Likelihood", "Mitigation", "Owner"], risk_rows)
    else:
        doc.add_paragraph("No risks identified.")
        doc.add_paragraph()

    # Budget Summary
    doc.add_heading("Budget Summary", level=1)
    doc.add_paragraph(data.get("budget_summary", "No budget information provided."))
    doc.add_paragraph()

    # Upcoming Activities
    doc.add_heading("Upcoming Activities", level=1)
    _add_list(doc, data.get("upcoming_activities", []))

    # Decisions Needed
    doc.add_heading("Decisions Needed", level=1)
    _add_list(doc, data.get("decisions_needed", []))

    # Footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Report generated on {_format_date(data.get('date', date.today()))}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _create_engineering_init_report(doc: "Document", data: Dict[str, Any]) -> None:
    """Create an engineering initiative report."""
    # Title
    title = doc.add_heading(data.get("title", "Engineering Initiative Report"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    _add_metadata_section(doc, data)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data.get("summary", "No summary provided."))
    doc.add_paragraph()

    # Sponsors
    doc.add_heading("Sponsors", level=1)
    _add_list(doc, data.get("sponsors", []))

    # Team
    doc.add_heading("Team", level=1)
    team_members = data.get("team_members", [])
    if team_members:
        team_rows = []
        for member in team_members:
            team_rows.append([
                member.get("name", "N/A"),
                member.get("role", "N/A"),
                member.get("email", "N/A")
            ])
        _add_table(doc, ["Name", "Role", "Email"], team_rows)
    else:
        doc.add_paragraph("No team members specified.")
        doc.add_paragraph()

    # Objectives
    doc.add_heading("Objectives", level=1)
    _add_list(doc, data.get("objectives", []))

    # Scope
    doc.add_heading("Scope", level=1)
    doc.add_paragraph(data.get("scope", "To be defined."))
    doc.add_paragraph()

    # Technical Details
    doc.add_heading("Technical Details", level=1)
    doc.add_paragraph(data.get("technical_details", "To be defined."))
    doc.add_paragraph()

    # Implementation Phases
    doc.add_heading("Implementation Phases", level=1)
    phases = data.get("implementation_phases", [])
    if phases:
        for i, phase in enumerate(phases, 1):
            phase_name = phase.get("name", f"Phase {i}")
            doc.add_heading(f"Phase {i}: {phase_name}", level=2)
            if phase.get("description"):
                doc.add_paragraph(phase["description"])
            if phase.get("duration"):
                p = doc.add_paragraph()
                p.add_run("Duration: ").bold = True
                p.add_run(str(phase["duration"]))
            if phase.get("deliverables"):
                p = doc.add_paragraph()
                p.add_run("Deliverables:").bold = True
                _add_list(doc, phase["deliverables"])
    else:
        doc.add_paragraph("No implementation phases defined.")
    doc.add_paragraph()

    # Success Criteria
    doc.add_heading("Success Criteria", level=1)
    _add_list(doc, data.get("success_criteria", []))

    # Milestones
    doc.add_heading("Milestones", level=1)
    milestones = data.get("milestones", [])
    if milestones:
        milestone_rows = []
        for milestone in milestones:
            status = milestone.get("status", {})
            status_value = status.get("value", "N/A") if isinstance(status, dict) else str(status)

            milestone_rows.append([
                milestone.get("name", "N/A"),
                _format_date(milestone.get("target_date")),
                status_value,
                f"{milestone.get('completion_percentage', 0)}%"
            ])
        _add_table(doc, ["Milestone", "Target Date", "Status", "Completion"], milestone_rows)
    else:
        doc.add_paragraph("No milestones defined.")
        doc.add_paragraph()

    # Impact Analysis
    doc.add_heading("Impact Analysis", level=1)
    doc.add_paragraph(data.get("impact_analysis", "To be defined."))
    doc.add_paragraph()

    # Resources Required
    doc.add_heading("Resources Required", level=1)
    doc.add_paragraph(data.get("resources_required", "To be defined."))
    doc.add_paragraph()

    # Risks & Mitigation
    doc.add_heading("Risks & Mitigation", level=1)
    risks = data.get("risks", [])
    if risks:
        for i, risk in enumerate(risks, 1):
            doc.add_heading(f"{i}. {risk.get('description', 'Risk')}", level=2)
            impact = risk.get("impact", {})
            likelihood = risk.get("likelihood", {})
            impact_value = impact.get("value", "N/A") if isinstance(impact, dict) else str(impact)
            likelihood_value = likelihood.get("value", "N/A") if isinstance(likelihood, dict) else str(likelihood)

            p = doc.add_paragraph()
            p.add_run("Impact: ").bold = True
            p.add_run(impact_value + "\n")
            p.add_run("Likelihood: ").bold = True
            p.add_run(likelihood_value + "\n")
            p.add_run("Mitigation: ").bold = True
            p.add_run(risk.get("mitigation", "N/A"))
            if risk.get("owner"):
                p.add_run("\nOwner: ").bold = True
                p.add_run(risk["owner"])
            doc.add_paragraph()
    else:
        doc.add_paragraph("No risks identified.")
        doc.add_paragraph()

    # Rollout Strategy
    doc.add_heading("Rollout Strategy", level=1)
    doc.add_paragraph(data.get("rollout_strategy", "To be defined."))
    doc.add_paragraph()

    # Footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Report generated on {_format_date(data.get('date', date.today()))}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
